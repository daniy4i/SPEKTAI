"""
Convert SPEKT-AI-PRD.html to SPEKT-AI-PRD.docx
"""

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

HTML_PATH  = "SPEKT-AI-PRD.html"
DOCX_PATH  = "SPEKT-AI-PRD.docx"

# Brand palette (light-mode friendly for Word)
VOID   = RGBColor(0x08, 0x08, 0x0D)
NEON   = RGBColor(0x2B, 0xCC, 0x0A)   # slightly darkened for white bg legibility
RED    = RGBColor(0xE9, 0x45, 0x60)
BLUE   = RGBColor(0x4A, 0x7C, 0xDB)
GOLD   = RGBColor(0xD4, 0xA8, 0x43)
GRAY   = RGBColor(0x55, 0x55, 0x66)
BLACK  = RGBColor(0x08, 0x08, 0x0D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_font(run, name="Georgia", size=11, bold=False, italic=False, color=None):
    run.font.name     = name
    run.font.size     = Pt(size)
    run.font.bold     = bold
    run.font.italic   = italic
    if color:
        run.font.color.rgb = color

def set_para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_section_divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 72)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ── Parse HTML ─────────────────────────────────────────────────────────────────
with open(HTML_PATH, "r", encoding="utf-8") as f:
    soup = BeautifulSoup(f.read(), "lxml")

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover page ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SPEKT AI")
set_font(r, "Georgia", 42, bold=False, color=BLACK)
set_para_spacing(p, before=24, after=4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Product Requirements Document  ·  v1.0")
set_font(r, "Courier New", 10, color=GRAY)
set_para_spacing(p, before=0, after=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Confidential — For Investor Review Only")
set_font(r, "Courier New", 9, color=RED)
set_para_spacing(p, before=0, after=32)

add_section_divider(doc)

# ── Walk the HTML body ─────────────────────────────────────────────────────────
# We pull content in document order, skipping <style>, <script>, <nav>

def clean_text(t):
    return re.sub(r'\s+', ' ', t).strip()

def process_node(node, doc):
    if not hasattr(node, 'name') or node.name is None:
        return  # text node handled by parent

    tag = node.name.lower()

    # Skip decorative / scripted elements
    if tag in ('style', 'script', 'nav', 'head', 'meta', 'link', 'canvas'):
        return

    # ── Headings ──────────────────────────────────────────────────────────────
    if tag == 'h1':
        text = clean_text(node.get_text())
        if not text:
            return
        add_section_divider(doc)
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, "Georgia", 28, color=BLACK)
        set_para_spacing(p, before=18, after=6)
        return

    if tag == 'h2':
        text = clean_text(node.get_text())
        if not text:
            return
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, "Georgia", 20, color=BLACK)
        set_para_spacing(p, before=14, after=4)
        return

    if tag == 'h3':
        text = clean_text(node.get_text())
        if not text:
            return
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, "Georgia", 14, bold=True, color=BLACK)
        set_para_spacing(p, before=10, after=3)
        return

    if tag == 'h4':
        text = clean_text(node.get_text())
        if not text:
            return
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        set_font(r, "Courier New", 9, bold=True, color=GRAY)
        set_para_spacing(p, before=8, after=2)
        return

    # ── Paragraphs ────────────────────────────────────────────────────────────
    if tag == 'p':
        text = clean_text(node.get_text())
        if not text:
            return
        classes = node.get('class', [])
        p = doc.add_paragraph()

        # Label style (Courier, uppercase, small)
        if 'label' in classes or 'overline' in classes:
            r = p.add_run(text.upper())
            set_font(r, "Courier New", 9, color=GRAY)
        elif 'stat-value' in classes or any('stat' in c for c in classes):
            r = p.add_run(text)
            set_font(r, "Georgia", 28, color=BLACK)
        else:
            r = p.add_run(text)
            set_font(r, "Georgia", 11, color=BLACK)
        set_para_spacing(p, before=0, after=5)
        return

    # ── Lists ─────────────────────────────────────────────────────────────────
    if tag in ('ul', 'ol'):
        for li in node.find_all('li', recursive=False):
            text = clean_text(li.get_text())
            if not text:
                continue
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(text)
            set_font(r, "Georgia", 11, color=BLACK)
            set_para_spacing(p, before=1, after=2)
        return

    # ── Tables ────────────────────────────────────────────────────────────────
    if tag == 'table':
        rows = node.find_all('tr')
        if not rows:
            return

        # Count max columns
        max_cols = max(
            sum(int(td.get('colspan', 1)) for td in row.find_all(['th', 'td']))
            for row in rows
        )
        if max_cols == 0:
            return

        tbl = doc.add_table(rows=0, cols=max_cols)
        tbl.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells_html = row.find_all(['th', 'td'])
            if not cells_html:
                continue
            row_cells = tbl.add_row().cells
            col_idx = 0
            for cell_html in cells_html:
                if col_idx >= max_cols:
                    break
                text = clean_text(cell_html.get_text())
                is_header = cell_html.name == 'th' or i == 0
                row_cells[col_idx].text = text
                p = row_cells[col_idx].paragraphs[0]
                if p.runs:
                    r = p.runs[0]
                    if is_header:
                        set_font(r, "Courier New", 8, bold=True, color=BLACK)
                        shade_cell(row_cells[col_idx], "E8E6E1")
                    else:
                        set_font(r, "Georgia", 10, color=BLACK)
                col_idx += int(cell_html.get('colspan', 1))

        doc.add_paragraph()
        return

    # ── Dividers ──────────────────────────────────────────────────────────────
    if tag == 'hr':
        add_section_divider(doc)
        return

    # ── Generic block containers — recurse into children ─────────────────────
    if tag in ('div', 'section', 'article', 'main', 'header', 'footer',
               'aside', 'figure', 'body', 'html'):
        for child in node.children:
            process_node(child, doc)
        return

    # ── Inline elements at top level (rare) ───────────────────────────────────
    if tag in ('span', 'strong', 'em', 'b', 'i', 'a', 'code', 'pre'):
        text = clean_text(node.get_text())
        if text:
            p = doc.add_paragraph()
            r = p.add_run(text)
            set_font(r, "Georgia", 11, color=BLACK)
            set_para_spacing(p, before=0, after=5)
        return

# Start from <body>
body = soup.find('body')
if body:
    for child in body.children:
        process_node(child, doc)

# ── Footer note ────────────────────────────────────────────────────────────────
add_section_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SPEKT AI  ·  Confidential  ·  2026  ·  ds@schotzenterprises.com")
set_font(r, "Courier New", 8, color=GRAY)

doc.save(DOCX_PATH)
print(f"Saved → {DOCX_PATH}")
